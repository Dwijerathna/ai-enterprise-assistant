"""File parsing utilities for supported document formats."""

from pathlib import Path
from typing import Any


class UnsupportedFileTypeError(ValueError):
    """Raised when a file type is not supported by the parser."""


def extract_text(file_path: str) -> str:
    """
    Extract plain text from a supported document file.

    Supported formats: PDF, TXT, DOCX.
    """
    pages = extract_document_pages(file_path)
    return "\n\n".join(page["text"] for page in pages if page["text"])


def extract_document_pages(file_path: str) -> list[dict[str, Any]]:
    """
    Extract page-aware text sections from a supported document file.

    Each page dict contains:
    - page_number
    - section_title
    - text
    """
    path = _validate_path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return [{"page_number": 1, "section_title": None, "text": _extract_txt(path)}]
    if suffix == ".pdf":
        return _extract_pdf_pages(path)
    if suffix == ".docx":
        return _extract_docx_pages(path)

    raise UnsupportedFileTypeError(
        f"Unsupported file type '{suffix}'. Supported types: .pdf, .txt, .docx"
    )


def _validate_path(file_path: str) -> Path:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if not path.is_file():
        raise ValueError(f"Path is not a file: {file_path}")
    return path


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_pdf(path: Path) -> str:
    pages = _extract_pdf_pages(path)
    return "\n\n".join(page["text"] for page in pages if page["text"])


def _extract_pdf_pages(path: Path) -> list[dict[str, Any]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[dict[str, Any]] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(
            {
                "page_number": index,
                "section_title": f"Page {index}",
                "text": text,
            }
        )
    return pages


def _extract_docx(path: Path) -> str:
    pages = _extract_docx_pages(path)
    return "\n\n".join(page["text"] for page in pages if page["text"])


def _extract_docx_pages(path: Path) -> list[dict[str, Any]]:
    from docx import Document

    document = Document(str(path))
    sections: list[dict[str, Any]] = []
    current_title: str | None = None
    current_lines: list[str] = []

    def flush_section() -> None:
        if not current_lines:
            return
        sections.append(
            {
                "page_number": 1,
                "section_title": current_title,
                "text": "\n".join(current_lines),
            }
        )

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name.lower().startswith("heading"):
            flush_section()
            current_lines = []
            current_title = text
            continue
        current_lines.append(text)

    flush_section()
    if sections:
        return sections

    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    return [{"page_number": 1, "section_title": None, "text": "\n".join(paragraphs)}]
